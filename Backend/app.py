import os
import re
import io
import json
import base64
import tempfile
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS

# PDF/Doc processing
import fitz  # PyMuPDF
import pdfplumber
from PIL import Image, ImageDraw, ImageFont
import openpyxl
from docx import Document

import easyocr
import numpy as np

# Initialize the reader once at the top of your file (outside the function)
reader = easyocr.Reader(['en'])

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = tempfile.mkdtemp()


PATTERNS = {
    "Aadhaar Number": {
        "regex": r"\b[2-9]\d{3}\s?\d{4}\s?\d{4}\b",
        "color": "#FF4444",
        "icon": "🪪",
        "severity": "critical"
    },
    "PAN Card": {
        "regex": r"\b[A-Z]{5}[0-9]{4}[A-Z]\b",
        "color": "#FF6B35",
        "icon": "💳",
        "severity": "critical"
    },
    "Phone Number": {
        "regex": r"(?:\+91[\-\s]?)?(?:\(?[6-9]\d{4}\)?[\-\s]?)?\d{5}[\-\s]?\d{5}\b|\b[6-9]\d{9}\b",
        "color": "#F59E0B",
        "icon": "📱",
        "severity": "high"
    },
    "Email Address": {
        "regex": r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b",
        "color": "#3B82F6",
        "icon": "📧",
        "severity": "high"
    },
    "Date of Birth": {
        "regex": r"\b(?:0?[1-9]|[12]\d|3[01])[-/.](?:0?[1-9]|1[0-2])[-/.](?:19|20)\d{2}\b",
        "color": "#8B5CF6",
        "icon": "🎂",
        "severity": "medium"
    },
    "Passport Number": {
        "regex": r"\b[A-PR-WYa-pr-wy][1-9]\d\s?\d{4}[1-9]\b",
        "color": "#EF4444",
        "icon": "📘",
        "severity": "critical"
    },
    "Driving Licence": {
        "regex": r"\b[A-Z]{2}[-\s]?\d{2}[-\s]?\d{4}[-\s]?\d{7}\b",
        "color": "#EC4899",
        "icon": "🚗",
        "severity": "high"
    },
    "Bank Account Number": {
        "regex": r"\b\d{9,18}\b(?=.*(?:account|acc|a/c|bank)|\s*\d{4})",
        "color": "#10B981",
        "icon": "🏦",
        "severity": "critical"
    },
    "IFSC Code": {
        "regex": r"\b[A-Z]{4}0[A-Z0-9]{6}\b",
        "color": "#06B6D4",
        "icon": "🏛️",
        "severity": "high"
    },
    "Credit/Debit Card": {
        "regex": r"\b(?:4\d{3}|5[1-5]\d{2}|6011|3[47]\d{2})[\s\-]?\d{4}[\s\-]?\d{4}[\s\-]?\d{4}\b",
        "color": "#F97316",
        "icon": "💰",
        "severity": "critical"
    },
    "GST Number": {
        "regex": r"\b\d{2}[A-Z]{5}\d{4}[A-Z][1-9A-Z]Z[0-9A-Z]\b",
        "color": "#84CC16",
        "icon": "🧾",
        "severity": "medium"
    },
    "IP Address": {
        "regex": r"\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\b",
        "color": "#6366F1",
        "icon": "🌐",
        "severity": "low"
    },
    "Pincode": {
        "regex": r"\b[1-9]\d{5}\b",
        "color": "#14B8A6",
        "icon": "📍",
        "severity": "low"
    },
    "Voter ID": {
        "regex": r"\b[A-Z]{3}\d{7}\b",
        "color": "#A855F7",
        "icon": "🗳️",
        "severity": "high"
    },
}


def scan_text(text):
    """Scan text and return list of findings."""
    findings = []
    for label, cfg in PATTERNS.items():
        for m in re.finditer(cfg["regex"], text, re.IGNORECASE):
            findings.append({
                "type": label,
                "value": m.group(),
                "start": m.start(),
                "end": m.end(),
                "color": cfg["color"],
                "icon": cfg["icon"],
                "severity": cfg["severity"],
            })
    # De-duplicate: keep longest match at same position
    findings.sort(key=lambda x: (x["start"], -(x["end"] - x["start"])))
    deduped, seen = [], set()
    for f in findings:
        overlap = any(not (f["end"] <= s or f["start"] >= e) for s, e in seen)
        if not overlap:
            deduped.append(f)
            seen.add((f["start"], f["end"]))
    return deduped


def mask_text(text, findings):
    """Replace sensitive values in text with black bars (█ chars)."""
    result = list(text)
    for f in sorted(findings, key=lambda x: x["start"], reverse=True):
        bar = "█" * len(f["value"])
        result[f["start"]:f["end"]] = list(bar)
    return "".join(result)



def redact_pdf(pdf_bytes, findings_by_page):
    """Use PyMuPDF to draw black redaction rectangles over sensitive text."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page_num, page in enumerate(doc):
        page_findings = findings_by_page.get(page_num, [])
        for item in page_findings:
            # Search for text instance on page
            areas = page.search_for(item["value"])
            for rect in areas:
                # Add redaction annotation
                page.add_redact_annot(rect, fill=(0, 0, 0))
        page.apply_redactions()
    out = io.BytesIO()
    doc.save(out)
    doc.close()
    return out.getvalue()


def extract_pdf_findings(pdf_bytes):
    """Extract text per page, scan each, return all findings + per-page map."""
    all_findings = []
    findings_by_page = {}
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    full_text = ""
    for page_num, page in enumerate(doc):
        text = page.get_text()
        full_text += f"\n[Page {page_num+1}]\n" + text
        page_findings = scan_text(text)
        if page_findings:
            findings_by_page[page_num] = page_findings
        all_findings.extend(page_findings)
    doc.close()
    
    seen_vals = set()
    unique = []
    for f in all_findings:
        if f["value"] not in seen_vals:
            unique.append(f)
            seen_vals.add(f["value"])
    return unique, findings_by_page, full_text



def process_image(img_bytes, filename):
    try:
        # 1. Convert bytes to PIL Image
        img_pil = Image.open(io.BytesIO(img_bytes)).convert("RGB")
        img_np = np.array(img_pil) # EasyOCR needs a numpy array
        
        # 2. Extract Text with EasyOCR
        # detail=0 returns just the strings, which is faster
        results = reader.readtext(img_np) 
        
        # results looks like: [([[x,y],[x,y]...], 'text', confidence), ...]
        full_text = " ".join([res[1] for res in results])
        findings = scan_text(full_text)
        
        # 3. Redact the Image
        draw = ImageDraw.Draw(img_pil)
        for res in results:
            box, text, conf = res
            # Check if this specific piece of text is part of any sensitive finding
            for f in findings:
                if text in f["value"] or f["value"] in text:
                    # box is [[x,y], [x,y], [x,y], [x,y]]
                    top_left = tuple(box[0])
                    bottom_right = tuple(box[2])
                    draw.rectangle([top_left, bottom_right], fill="black")
        
        # 4. Encode to Base64
        buf = io.BytesIO()
        img_pil.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode()
        
        return findings, full_text, b64
    except Exception as e:
        print(f"Image processing error: {e}")
        return [], f"Error: {str(e)}", None


def process_excel(excel_bytes):
    wb = openpyxl.load_workbook(io.BytesIO(excel_bytes))
    all_text = ""
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value:
                    all_text += str(cell.value) + " "
    findings = scan_text(all_text)
    # Redact
    wb2 = openpyxl.load_workbook(io.BytesIO(excel_bytes))
    for sheet in wb2.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value:
                    val = str(cell.value)
                    cell_findings = scan_text(val)
                    if cell_findings:
                        cell.value = mask_text(val, cell_findings)
    out = io.BytesIO()
    wb2.save(out)
    return findings, all_text, out.getvalue()


def process_docx(docx_bytes):
    doc = Document(io.BytesIO(docx_bytes))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    findings = scan_text(all_text)
    # Redact
    doc2 = Document(io.BytesIO(docx_bytes))
    for para in doc2.paragraphs:
        if para.text:
            pf = scan_text(para.text)
            if pf:
                masked = mask_text(para.text, pf)
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = masked
    out = io.BytesIO()
    doc2.save(out)
    return findings, all_text, out.getvalue()


def process_txt(txt_bytes):
    text = txt_bytes.decode("utf-8", errors="replace")
    findings = scan_text(text)
    masked = mask_text(text, findings)
    return findings, text, masked.encode("utf-8")


@app.route("/api/analyze", methods=["POST"])
def analyze():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    f = request.files["file"]
    filename = f.filename.lower()
    raw = f.read()
    ext = filename.rsplit(".", 1)[-1] if "." in filename else ""

    result = {
        "filename": f.filename,
        "file_type": ext.upper(),
        "findings": [],
        "summary": {},
        "redacted_b64": None,
        "redacted_ext": ext,
        "raw_text_preview": "",
        "page_count": 1,
    }

    try:
        if ext == "pdf":
            findings, findings_by_page, full_text = extract_pdf_findings(raw)
            redacted_bytes = redact_pdf(raw, findings_by_page)
            result["findings"] = findings
            result["raw_text_preview"] = full_text[:3000]
            result["redacted_b64"] = base64.b64encode(redacted_bytes).decode()
            result["page_count"] = len(fitz.open(stream=raw, filetype="pdf"))

        elif ext in ("png", "jpg", "jpeg", "webp", "bmp", "tiff"):
            findings, text, img_b64 = process_image(raw, filename)
            result["findings"] = findings
            result["raw_text_preview"] = text[:3000]
            result["redacted_b64"] = img_b64
            result["redacted_ext"] = "png"

        elif ext in ("xlsx", "xls"):
            findings, text, redacted_bytes = process_excel(raw)
            result["findings"] = findings
            result["raw_text_preview"] = text[:3000]
            result["redacted_b64"] = base64.b64encode(redacted_bytes).decode()

        elif ext == "docx":
            findings, text, redacted_bytes = process_docx(raw)
            result["findings"] = findings
            result["raw_text_preview"] = text[:3000]
            result["redacted_b64"] = base64.b64encode(redacted_bytes).decode()

        elif ext == "txt":
            findings, text, redacted_bytes = process_txt(raw)
            result["findings"] = findings
            result["raw_text_preview"] = text[:3000]
            result["redacted_b64"] = base64.b64encode(redacted_bytes).decode()

        else:
           
            try:
                text = raw.decode("utf-8", errors="replace")
                findings = scan_text(text)
                masked = mask_text(text, findings).encode("utf-8")
                result["findings"] = findings
                result["raw_text_preview"] = text[:3000]
                result["redacted_b64"] = base64.b64encode(masked).decode()
            except Exception:
                return jsonify({"error": f"Unsupported file type: {ext}"}), 400

        severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0}
        type_counts = {}
        unique_values = set()
        for item in result["findings"]:
            sev = item.get("severity", "low")
            severity_counts[sev] = severity_counts.get(sev, 0) + 1
            t = item["type"]
            type_counts[t] = type_counts.get(t, 0) + 1
            unique_values.add(item["value"])

        result["summary"] = {
            "total_findings": len(result["findings"]),
            "unique_values": len(unique_values),
            "severity_counts": severity_counts,
            "type_counts": type_counts,
            "risk_score": min(100, severity_counts["critical"] * 25 + severity_counts["high"] * 15 + severity_counts["medium"] * 8 + severity_counts["low"] * 3),
        }

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    return jsonify(result)


@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "version": "1.0.0"})


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
